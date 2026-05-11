from docx import Document

doc = Document("attached_assets/ВЗИВ_08_05_Диплом_Ксёнжик_Александр_2220_1778488331470.docx")

print("=== Paragraphs containing ИСПРАВИТЬ or СТИЛЬ ===")
for i, p in enumerate(doc.paragraphs):
    t = p.text
    if "ИСПРАВИТЬ" in t or "СТИЛЬ" in t.upper() and "СТИЛ" in t or t.endswith("СТИЛЬ!!!.") or "СТИЛЬ!!!" in t:
        print(f"[{i:4d}] {t}")
        print()

print()
print("=== References section (from СПИСОК) ===")
in_refs = False
for i, p in enumerate(doc.paragraphs):
    if "СПИСОК ИСПОЛЬЗ" in p.text and not p.style.name.startswith("toc"):
        in_refs = True
    if in_refs:
        print(f"[{i:4d}] ({p.style.name}) {p.text[:200]}")

print()
print("=== Section 2.1 paragraphs (where XP definition could go) ===")
for i in range(278, 310):
    p = doc.paragraphs[i]
    print(f"[{i:4d}] ({p.style.name[:14]}) {p.text[:180]}")

print()
print("=== Section 2.2 paragraphs (where CMS rationale should go) ===")
for i in range(305, 360):
    p = doc.paragraphs[i]
    print(f"[{i:4d}] ({p.style.name[:14]}) {p.text[:180]}")
