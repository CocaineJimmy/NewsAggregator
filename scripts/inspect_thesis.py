from docx import Document
import sys

src = "attached_assets/ВЗИВ_08_05_Диплом_Ксёнжик_Александр_2220_1778488331470.docx"
doc = Document(src)

print(f"Total paragraphs: {len(doc.paragraphs)}")
print(f"Total tables: {len(doc.tables)}")
print(f"Sections: {len(doc.sections)}")

# Print all paragraphs with index, style, and short text
markers = ["ИСПРАВИТЬ", "СТИЛЬ", "XP", "Рисунок 2", "СПИСОК ИСПОЛЬЗ", "РЕФЕРАТ",
           "2.1.1", "2.2.1", "2.4.1", "2.4 ", "2.4.4", "2.5.4", "2.5.7", "2.6.1",
           "2.7", "2.8", "ОГЛАВЛЕНИЕ", "ВВЕДЕНИЕ", "ЗАКЛЮЧЕНИЕ", "ГЛАВА 1", "ГЛАВА 2", "ГЛАВА 3"]

for i, p in enumerate(doc.paragraphs):
    t = p.text.strip()
    if any(m in t for m in markers) or (t and i < 30):
        style = p.style.name if p.style else "?"
        print(f"[{i:4d}] ({style[:14]:14s}) {t[:160]}")
